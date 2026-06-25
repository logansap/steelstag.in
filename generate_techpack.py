from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

# ── Helpers ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        val = kwargs.get(side, {'sz':'4','val':'single','color':'CCCCCC'})
        el  = OxmlElement(f'w:{side}')
        for k,v in val.items():
            el.set(qn(f'w:{k}'), v)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def heading(text, size=13, bold=True, color='1A1A1A', space_before=14, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold      = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return p

def subheading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text.upper())
    run.bold      = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string('6B7280')
    return p

def body(text, size=9.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string('1A1A1A')
    return p

def kv(key, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    k = p.add_run(f'{key}:  ')
    k.bold = True
    k.font.size = Pt(9.5)
    k.font.color.rgb = RGBColor.from_string('374151')
    v = p.add_run(value)
    v.font.size = Pt(9.5)
    v.font.color.rgb = RGBColor.from_string('1A1A1A')
    return p

def make_table(headers, rows, header_bg='1A1A1A', header_fg='FFFFFF', alt_bg='F9FAFB'):
    col_count = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=col_count)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(header_fg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        bg  = alt_bg if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(8.5)
            if ci == 0:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string('374151')
            else:
                run.font.color.rgb = RGBColor.from_string('1A1A1A')
    return t

def hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E5E7EB')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ════════════════════════════════════════════════════════════════════════════
# COVER / HEADER
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(2)
run = p.add_run('STEELSTAG')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor.from_string('1A1A1A')

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(0)
run2 = p2.add_run('TECHNICAL SPECIFICATION PACKAGE  ·  PREMIUM SOLID TEES  ·  SS2026')
run2.font.size = Pt(8)
run2.font.color.rgb = RGBColor.from_string('6B7280')

hr()

# ════════════════════════════════════════════════════════════════════════════
# 1. BRAND & ORDER DETAILS
# ════════════════════════════════════════════════════════════════════════════
heading('1.  BRAND & ORDER DETAILS')
kv('Brand',           'SteelStag')
kv('Contact',         'hello@steelstag.in')
kv('Season',          'SS2026 — Initial Launch Run')
kv('Styles',          'ST-RN-001 (Round Neck)  |  ST-DS-002 (Deep Scoop)  |  ST-VN-003 (V-Neck)')
kv('Sizes',           'S  /  M  /  L  /  XL')
kv('Size Split',      'S 15%  |  M 35%  |  L 35%  |  XL 15%')
kv('Destination',     'India — Direct to Consumer (DTC)')

hr()

# ════════════════════════════════════════════════════════════════════════════
# 2. FABRIC SPECIFICATION
# ════════════════════════════════════════════════════════════════════════════
heading('2.  FABRIC SPECIFICATION')
kv('Fabric',          '100% Combed Cotton — Single Jersey')
kv('Yarn Count',      '30s Combed Ring Spun')
kv('GSM',             '125 GSM  (tolerance ±5 GSM)')
kv('Construction',    'Single Jersey')
kv('Dyeing',          'Reactive Dyed')
kv('Finish',          'Light enzyme wash (no heavy silicone finish)')
kv('Pre-treatment',   'Pre-washed, pre-shrunk fabric')
kv('Shrinkage',       'Max ±3% after wash (no twisting)')

hr()

# ════════════════════════════════════════════════════════════════════════════
# 3. COLOR PALETTE
# ════════════════════════════════════════════════════════════════════════════
heading('3.  COLOR PALETTE')
body('All colors to be approved on pre-production lab dip before bulk. Pantone TCX references below.', size=8.5)
doc.add_paragraph()

color_headers = ['#', 'Color Name', 'Pantone TCX', 'Style Code', 'Notes']
color_rows = [
    ['01', 'Jet Black',      '19-0303 TCX', 'All styles', 'Priority — must have'],
    ['02', 'Bright White',   '11-0601 TCX', 'All styles', 'Priority — must have'],
    ['03', 'Seaport Teal',   '17-5126 TCX', 'All styles', 'Priority — hero color'],
    ['04', 'Wild Ginger',    '15-1157 TCX', 'All styles', 'Priority — hero color'],
    ['05', 'Brown Sugar',    '18-1048 TCX', 'All styles', ''],
]
make_table(color_headers, color_rows)
body('Note: 5 colors across all 3 styles. Provide lab dips for all colors. Priority colors (01–04) to be sampled first.', size=8)

hr()

# ════════════════════════════════════════════════════════════════════════════
# 4. CONSTRUCTION & STITCH SPEC
# ════════════════════════════════════════════════════════════════════════════
heading('4.  CONSTRUCTION & STITCH SPECIFICATION')

subheading('General')
kv('Stitch Density',  '12–14 SPI (stitches per inch)')
kv('Thread',          'Self-colour polyester thread to match each colorway')
kv('Seam Type',       'Overlock (serger) — all body seams')
kv('Side Seam',       'Clean overlock — no twisting')

subheading('Neck Finish')
body('All styles: Clean-bound neck finish — no rib band. Self-fabric tape, folded and stitched flat. Single needle topstitch 1/8" from edge, self-colour thread. Neck must lie flat against body without rolling or gaping. Reference: white tee reference image provided.', size=8.5)
doc.add_paragraph()
make_table(
    ['Style', 'Neck Construction', 'Detail'],
    [
        ['ST-RN-001  Round Neck', 'Clean-bound / self-fabric tape', 'No rib band — folded self-fabric, single needle topstitch 1/8" from edge, lies flat'],
        ['ST-DS-002  Deep Scoop', 'Clean-bound / self-fabric tape', 'No rib band — folded self-fabric, single needle topstitch 1/8" from edge, lies flat'],
        ['ST-VN-003  V-Neck',     'Clean-bound / self-fabric tape', 'No rib band — mitered V-point, folded self-fabric, single needle topstitch 1/8" from edge'],
    ]
)

subheading('Sleeve Finish')
make_table(
    ['Style', 'Sleeve Stitch', 'Detail'],
    [
        ['ST-RN-001  Round Neck', 'Clean flat hem', 'Twin needle, 1/4" hem, self-colour'],
        ['ST-DS-002  Deep Scoop', 'Clean flat hem', 'Twin needle, 1/4" hem, self-colour'],
        ['ST-VN-003  V-Neck',     'Clean flat hem', 'Twin needle, 1/4" hem, self-colour'],
    ]
)

subheading('Hem')
kv('Bottom Hem',      'Single needle, 1" hem, self-colour thread')
kv('Tolerance',       '±0.5" on all measurements')

hr()

# ════════════════════════════════════════════════════════════════════════════
# 5. MEASUREMENTS — BODY
# ════════════════════════════════════════════════════════════════════════════
heading('5.  BODY MEASUREMENTS  (inches)')
body('All measurements in inches. Tolerance ±0.5" unless stated.', size=8.5)
doc.add_paragraph()

body_headers = ['Point of Measure', 'S', 'M', 'L', 'XL']
body_rows = [
    ['HPS Length',              '26¾',  '27½',  '28¼',  '29'],
    ['CB Length (from seam)',   '25¾',  '26½',  '27¼',  '28'],
    ['Shoulder to Shoulder',    '17',   '17½',  '18',   '18½'],
    ['½ Chest (1" below AH)',   '19¼',  '19¾',  '20¼',  '20¾'],
    ['½ Waist',                 '19¼',  '19¾',  '20¼',  '20¾'],
    ['½ Sweep (Straight)',      '19¼',  '19¾',  '20¼',  '20¾'],
    ['Armhole Drop from HPS',   '10½',  '10¾',  '11',   '11¼'],
    ['½ Sleeve Opening',        '6',    '6¼',   '6½',   '6¾'],
]
make_table(body_headers, body_rows)

hr()

# ════════════════════════════════════════════════════════════════════════════
# 6. MEASUREMENTS — NECK
# ════════════════════════════════════════════════════════════════════════════
heading('6.  NECK MEASUREMENTS  (inches)')
doc.add_paragraph()

subheading('Round Neck — ST-RN-001')
make_table(
    ['Point of Measure', 'S', 'M', 'L', 'XL'],
    [
        ['Neck Width',       '7',     '7¼',    '7½',    '7¾'],
        ['Front Neck Drop',  '2⅞',    '3',     '3⅛',    '3¼'],
        ['Back Neck Drop',   '1',     '1',     '1',     '1'],
        ['Collar Band Ht',   'NA',    'NA',    'NA',    'NA'],
    ]
)

doc.add_paragraph()
subheading('Deep Scoop — ST-DS-002')
make_table(
    ['Point of Measure', 'S', 'M', 'L', 'XL'],
    [
        ['Neck Width',       '8',     '8½',    '8¾',    '9'],
        ['Front Neck Drop',  '4¼',    '4½',    '4⅝',    '4¾'],
        ['Back Neck Drop',   '1',     '1',     '1',     '1'],
        ['Collar Band Ht',   'NA',    'NA',    'NA',    'NA'],
    ]
)

doc.add_paragraph()
subheading('V-Neck — ST-VN-003')
make_table(
    ['Point of Measure', 'S', 'M', 'L', 'XL'],
    [
        ['Neck Width',       '6¾',    '7',     '7¼',    '7½'],
        ['Front Neck Drop',  '3½',    '3¾',    '4',     '4¼'],
        ['Back Neck Drop',   '1',     '1',     '1',     '1'],
        ['V-Point Depth',    '3½',    '3¾',    '4',     '4¼'],
        ['Collar Band Ht',   'NA',    'NA',    'NA',    'NA'],
    ]
)

hr()

# ════════════════════════════════════════════════════════════════════════════
# 7. LABEL & BRANDING PLACEMENT
# ════════════════════════════════════════════════════════════════════════════
heading('7.  LABEL & BRANDING PLACEMENT')
make_table(
    ['Label', 'Type', 'Position', 'Notes'],
    [
        ['Neck Label',      'Woven / printed',   'Centre back, 1" below collar seam',     'Per size: S / M / L / XL — files attached'],
        ['Wash Care Label', 'Woven / printed',   'Left side seam, 3" from bottom hem',    'File attached'],
        ['Brand Tag',       'Hang tag',          'Attached to left side seam with thread', 'Front + Back — files attached'],
        ['Size Sticker',    'Printed sticker',   'On polybag outside',                    'Per size — files attached'],
    ]
)
body('All label artwork files (.png, 300 DPI) attached separately. Do not substitute fonts or colors.', size=8)

hr()

# ════════════════════════════════════════════════════════════════════════════
# 8. PACKAGING
# ════════════════════════════════════════════════════════════════════════════
heading('8.  PACKAGING')

subheading('Primary Packaging — Kraft Paper Standup Pouch')
kv('Type',            'Kraft paper standup pouch with zipper seal')
kv('Size',            '32cm × 40cm (fits folded tee — standard retail fold)')
kv('Material',        'Kraft paper laminated, food-grade inner layer not required')
kv('Design',          'Two-tone: top 45% natural kraft (#C8A96E), bottom 55% matte black (#1A1A1A)')
kv('Logo Print',      '"STEELSTAG" centered on kraft section — large, black ink, serif or brand font')
kv('Tagline',         '"Color. Nothing else." below brand name — smaller, black ink')
kv('URL',             '"steelstag.in" printed small on dark top band — gold/kraft ink')
kv('Closure',         'Zipper reseal at top — allows customer to reuse pouch')
kv('Finish',          'Matte throughout — no gloss lamination')
kv('Reference',       'Similar to Finn Design (finn-design.de) packaging style')

subheading('Folding & Packing')
kv('Folding',         'Standard retail fold — tee folded to approx 28cm × 20cm before pouching')
kv('Insert',          'Size sticker (provided) affixed on outside of pouch — bottom right corner')
kv('Tissue',          'Optional: single sheet natural kraft tissue wrap inside pouch')

subheading('Shipping Carton')
kv('Carton',          'Export carton — size-wise packing per color')
kv('Carton Label',    'Brand: SteelStag | Style Code | Color | Size | Qty | PO No.')

hr()

# ════════════════════════════════════════════════════════════════════════════
# 9. SAMPLING & APPROVAL
# ════════════════════════════════════════════════════════════════════════════
heading('9.  SAMPLING & APPROVAL PROCESS')
make_table(
    ['Stage', 'Requirement', 'Approval'],
    [
        ['Lab Dip',          'All 5 colors — reactive dyed on approved 125 GSM combed cotton', 'Written approval required before bulk'],
        ['Fit Sample',       '1 pc per style in size M (Seaport Teal or Bright White)',        'Measurement + visual sign-off'],
        ['Pre-Production',   '1 pc per style × all 5 colors',                                  'Final approval before cutting'],
        ['Top of Production','2 pcs per style × per color from first bulk roll',     'Retained as bulk standard'],
    ]
)

hr()

# ════════════════════════════════════════════════════════════════════════════
# 10. QUALITY & COMPLIANCE
# ════════════════════════════════════════════════════════════════════════════
heading('10.  QUALITY & COMPLIANCE')
kv('Shrinkage',       'Max ±3% (length & width) after 3 washes at 30°C')
kv('Color Fastness',  'Min Grade 4 — washing, rubbing, perspiration')
kv('Pilling',         'Min Grade 3–4 after 1000 cycles')
kv('GSM Tolerance',   '±5 GSM from approved standard')
kv('AQL',             '2.5 — inline and final inspection')
kv('Certification',   'OCS / BCI certificate for cotton preferred — provide mill test report with sample')

hr()

# ════════════════════════════════════════════════════════════════════════════
# FOOTER
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
run = p.add_run('SteelStag  ·  hello@steelstag.in  ·  steelstag.in  ·  © 2026 SteelStag — Confidential')
run.font.size = Pt(7.5)
run.font.color.rgb = RGBColor.from_string('9CA3AF')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save
out = '/Users/I328057/workspace/Anaishu/SteelStag-TechPack-SS2026.docx'
doc.save(out)
print(f'Saved: {out}')
